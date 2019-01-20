import xlrd, xlwt, openpyxl, os
import sys_tools
from docx import Document
from modals.report_test_result import result_report
from common import settings
from docx.shared import Inches
from bs4 import BeautifulSoup


def get_report_id(report_name):
    report_id = 'NULL'
    for i in get_reports_IP_and_OP():
        temp_name = i.report_id[0:4] + i.report_name
        if report_name in temp_name:
            report_id = i.report_id

    return report_id


def get_report_module(report_id):
    for i in get_reports_IP_and_OP():
        module = 'Unknown'
        if report_id in i.report_id:
            module = i.report_module
            break
        else:
            pass
    return module


def get_log(report_name):
    if 'CC/MCC' in report_name:
        report_name = 'Top 50 CC_MCC Diagnoses'

    log_path = sys_tools.base_path + '/auto_results/logs/log.txt'
    log_path_target = sys_tools.base_path + '/auto_results/logs/' + report_name + '.txt'

    f = open(log_path, 'r')
    f_report = open(log_path_target, 'w')

    lines = f.readlines()
    steps_ = ''
    steps_report = []
    for i in lines:
        steps_ = steps_ + i
        steps_report.append(i)

    f_report.writelines(steps_report)
    f_report.close()
    f.close()

    f = open(log_path, 'w')
    f.write('')
    f.close()
    return steps_


def get_logs():
    log_path = sys_tools.base_path + '/auto_results/logs/log.txt'
    log_path_target = sys_tools.base_path + '/auto_results/logs/'
    f = open(log_path, 'r')

    lines = f.readlines()
    steps_ = {}
    steps = []
    j = 0
    for i in lines:
        steps.append(i)
        if 'report_names:>>>>>' in i:
            report_name = i.strip()[18:].rstrip('#')
            if 'CC/MCC' in report_name:
                report_name = 'Top 50 CC_MCC Diagnoses'
            j += 1
        if 'close browser' in i:
            f_report = open(log_path_target + report_name + '.txt', 'w')
            f_report.writelines(steps)
            f_report.close()
            steps_[report_name] = steps
            steps.clear()
    return steps_


def save_log(str_):
    if settings.log_print_switch:
        print(str_)

    log_path = sys_tools.base_path + '/auto_results/logs/log.txt'
    f = open(log_path, 'a')
    f.write(str_ + '\n')
    f.close()


def save_html_resource(page_resourece, report):
    file_path = sys_tools.base_path + '/auto_results/report_view_html_resources/' + report + '.html'
    f = open(file_path, 'w')
    f.write(page_resourece)
    f.close()


def clear_previous_auto_result():
    pass


def get_reports_IP_and_OP():
    f = sys_tools.base_path + '/resources/Legacy & Enhanced Report.xlsx'
    book = xlrd.open_workbook(f)  # 打开一个excel

    reports = []

    sheet_ip = book.sheet_by_name('IP')
    ip_nrow = sheet_ip.nrows

    sheet_op = book.sheet_by_name('OP')
    op_nrow = sheet_op.nrows

    for i in range(1, ip_nrow):
        report_id = sheet_ip.cell(i, 0).value.strip()
        report_module = sheet_ip.cell(i, 1).value.strip()
        report_name = sheet_ip.cell(i, 2).value.strip()
        report_name_saved_search = sheet_ip.cell(i, 3).value.strip()

        result = result_report(
            report_id=report_id,
            report_module=report_module,
            report_name=report_name,
            report_name_saved_search=report_name_saved_search,
        )
        reports.append(result)

    for i in range(1, op_nrow):
        report_id = sheet_op.cell(i, 0).value.strip()
        report_module = sheet_op.cell(i, 1).value.strip()
        report_name = sheet_op.cell(i, 2).value.strip()
        report_name_saved_search = sheet_op.cell(i, 3).value.strip()

        result = result_report(
            report_id=report_id,
            report_module=report_module,
            report_name=report_name,
            report_name_saved_search=report_name_saved_search,
        )
        reports.append(result)

    return reports


def save_reports_screenshot_as_html(result_report):
    file_path_target = sys_tools.base_path + '/auto_results/test_results/Test Results Screenshots.html'
    Fobj = open(file_path_target)
    Data = Fobj.read()
    Fobj.close()
    soup_elements = BeautifulSoup(Data, features="lxml")

    # get container_div
    container_div = soup_elements.select_one('#container')

    # create img div
    attrs_div = {'class': result_report.report_test_result}
    screenshot_div = soup_elements.new_tag('div', attrs=attrs_div)

    # create img title
    screenshot_title_p = soup_elements.new_tag('p')
    screenshot_title_p.string = result_report.report_name

    steps_attri = {
        'style': 'display:none',
    }
    screenshot_steps_p = soup_elements.new_tag('p', attrs=steps_attri)
    screenshot_steps_p.string = result_report.steps

    # create img tag
    attrs = {
        'class': result_report.report_test_result.__str__() + ' screenshot',
        'id': result_report.report_name,
        'alt': result_report.report_name,
        'src': '../../auto_results/screenshots/' + result_report.screen_shot
        # 'src': sys_tools.base_path + '\\auto_results\\screenshots\\' + result_report.screen_shot
    }
    screenshot_img = soup_elements.new_tag('img', attrs=attrs)

    # insert into container
    screenshot_div.insert(new_child=screenshot_img, position=2)
    screenshot_div.insert(new_child=screenshot_steps_p, position=1)
    screenshot_div.insert(new_child=screenshot_title_p, position=0)

    container_div.insert(new_child=screenshot_div, position=0)

    # save file
    file = open(file_path_target, 'w')
    file.write(str(soup_elements))
    file.close()


def write_test_result_report_word(result_report):

    image_path = sys_tools.base_path + '\\auto_results\\screenshots\\' + result_report.screen_shot
    word_result_path = sys_tools.base_path + '\\auto_results\\test_results\\' + settings.test_result_file_name_word
    # image_path = '../auto_results/screenshots/' + result_report.screen_shot
    # word_result_path = '../auto_results/test_results/' + settings.test_result_file_name_word

    if os.path.exists(word_result_path):
        document = Document(word_result_path)
    else:
        document = Document()
    document.add_heading(result_report.report_name + ': ' + result_report.report_test_result.__str__(), level=1)
    document.add_paragraph(text='Test Result: ' + result_report.msg)
    document.add_picture(image_path, width=Inches(4.5))
    document.save(word_result_path)


def write_test_result_report_excel(result_report):
    test_result_report_path = sys_tools.base_path + '\\auto_results\\test_results\\legacy reports auto test result.xlsx'
    path_screen_shot_path = sys_tools.base_path + '\\auto_results\screenshots\\' + result_report.screen_shot

    # test_result_report_path = '../../auto_results/test_results/legacy reports auto test result.xlsx'
    # path_screen_shot_path = '../../auto_results\screenshots/' + result_report.screen_shot

    wb = openpyxl.load_workbook(test_result_report_path)
    if 'IP' in result_report.report_id:
        ip_sheet = wb['IP']

        # insert test result
        ip_sheet.append([
            result_report.report_id,
            result_report.report_module,
            result_report.report_name,
            result_report.report_name_saved_search,
            result_report.report_test_result,
        ])

        # screenshot 设置超链接
        row_max = ip_sheet.max_row  # 获取最后一行
        ip_sheet.cell(row=row_max, column=6).hyperlink = path_screen_shot_path
        ip_sheet.cell(row=row_max, column=6).style = 'Hyperlink'
        ip_sheet.cell(row=row_max, column=6).value = result_report.screen_shot

        ip_sheet.cell(row=row_max, column=7).value = result_report.msg
    elif 'OP' in result_report.report_id:
        op_sheet = wb['OP']

        # insert test result
        op_sheet.append([
            result_report.report_id,
            result_report.report_module,
            result_report.report_name,
            result_report.report_name_saved_search,
            result_report.report_test_result,
        ])

        # screenshot 设置超链接
        row_max = op_sheet.max_row  # 获取最后一行
        op_sheet.cell(row=row_max, column=6).hyperlink = path_screen_shot_path
        op_sheet.cell(row=row_max, column=6).style = 'Hyperlink'
        op_sheet.cell(row=row_max, column=6).value = result_report.screen_shot

        op_sheet.cell(row=row_max, column=7).value = result_report.msg

    try:
        wb.save(test_result_report_path)
    except Exception as e:
        print('------write_test_result_report failed-------')
        print(e)


def write_test_result_com():
    test_result_report = result_report()

    book = xlwt.Workbook(encoding='utf-8', style_compression=0)

    sheet = book.add_sheet('IP', cell_overwrite_ok=True)

    book.save(sys_tools.base_path + '\legacy reports.xlsx')
